import os
import pdfplumber
import docx
import openpyxl
from PIL import Image
import pytesseract

# OCR ayarları (Django settings'ten; Django yüklü değilse env/varsayılan)
def _get_ocr_config():
    try:
        from django.conf import settings
        cmd = getattr(settings, "TESSERACT_CMD", None)
        lang = getattr(settings, "OCR_LANG", "tur")
        psm = getattr(settings, "OCR_PSM", "6")
    except Exception:
        cmd = os.environ.get("TESSERACT_CMD")
        lang = os.environ.get("OCR_LANG", "tur")
        psm = os.environ.get("OCR_PSM", "6")
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd
    return lang, psm


def extract_text_from_file(file_path):
    if not os.path.exists(file_path):
        return "Dosya bulunamadı"

    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            return extract_pdf(file_path)

        elif ext == ".docx":
            return extract_docx(file_path)

        elif ext == ".xlsx":
            return extract_excel(file_path)

        elif ext in [".png", ".jpg", ".jpeg"]:
            return extract_image(file_path)

        else:
            return "Desteklenmeyen dosya formatı"

    except Exception as e:
        return f"Hata oluştu: {str(e)}"


def extract_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_docx(path):
    """Metin + tabloları çıkarır. Birim Fiyat Teklif Cetveli gibi tablolar satır satır (sütunlar TAB ile) döner."""
    doc = docx.Document(path)
    parts = []
    # Önce tabloları al (cetvel genelde tablo formatındadır)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join((cell.text or "").strip().replace("\n", " ") for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
        if parts:
            parts.append("")  # tablo sonu
    # Paragrafları ekle (başlık, açıklama vb.)
    para_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    if para_text:
        parts.append(para_text)
    return "\n".join(parts) if parts else ""


def extract_excel(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    text = ""

    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join(str(cell) for cell in row if cell is not None)
            if row_text:
                text += row_text + "\n"

    return text


def extract_image(path):
    lang, psm = _get_ocr_config()
    img = Image.open(path)
    return pytesseract.image_to_string(
        img,
        lang=lang,
        config=f"--psm {psm}"
    )
